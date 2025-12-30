from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Literal
import uuid
from datetime import datetime, timezone
import bcrypt
import jwt
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
import io
import re
from docx import Document
from docx.shared import Pt, Inches
from emergentintegrations.llm.chat import LlmChat, UserMessage
import asyncio

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME', 'blum_treuhand')]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'blum-treuhand-secret-key-2025')
JWT_ALGORITHM = "HS256"
JWT_EXPIRATION_HOURS = 24

# OpenAI Configuration
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')

# Create the main app
app = FastAPI(title="Blum Verwaltungs- und Treuhand AG", version="1.0.0")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# ==================== MODELS ====================

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    role: Literal["admin", "user"] = "user"

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    email: str
    name: str
    role: str
    created_at: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserResponse

class PartyCreate(BaseModel):
    party_type: Literal["person", "firma"] = "person"
    name: str
    address: str = ""
    plz: str = ""
    ort: str = ""
    land: str = "Schweiz"
    email: str = ""
    phone: str = ""
    identification: str = ""
    role: Literal["aktionaer", "vr", "beide"] = "aktionaer"
    requires_signature: bool = True

class CompanyCreate(BaseModel):
    name_current: str
    name_new: Optional[str] = None
    sitz_current: str = ""
    sitz_new: Optional[str] = None
    zweck_current: str = ""
    zweck_new: Optional[str] = None
    che_nummer: str = ""
    hr_nummer: str = ""
    gruendungsdatum: Optional[str] = None
    aktienkapital: float = 0
    liberierung: Literal["voll", "teil"] = "voll"
    anzahl_aktien: int = 0
    nennwert: float = 0
    bankkonto: str = ""
    saldo: float = 0
    gewaehrleistung: bool = False
    gewaehrleistung_notiz: str = ""

class VRMemberCreate(BaseModel):
    name: str
    zeichnungsart: Literal["einzeln", "kollektiv"] = "einzeln"
    zeichnungsdetails: str = ""
    unterschreibt: bool = True

class DealCreate(BaseModel):
    deal_type: Literal["ankauf", "verkauf"]
    company: CompanyCreate
    sellers: List[PartyCreate] = []
    buyers: List[PartyCreate] = []
    vr_members: List[VRMemberCreate] = []
    kaufpreis: float = 0
    kaufpreis_regelung: str = ""
    besitzantritt: Optional[str] = None
    unterschriftsort: str = "Zürich"
    unterschriftsdatum: Optional[str] = None
    anzahlung: float = 0
    anzahlung_aktiviert: bool = False
    darlehen_uebernahme: bool = False
    darlehen_betrag: float = 0
    notizen: str = ""

class DealResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    deal_number: str
    deal_type: str
    status: str
    company: dict
    sellers: List[dict]
    buyers: List[dict]
    vr_members: List[dict]
    kaufpreis: float
    kaufpreis_regelung: str
    besitzantritt: Optional[str]
    unterschriftsort: str
    unterschriftsdatum: Optional[str]
    anzahlung: float
    anzahlung_aktiviert: bool
    darlehen_uebernahme: bool
    darlehen_betrag: float
    notizen: str
    created_at: str
    updated_at: str
    created_by: str

class DealStatusUpdate(BaseModel):
    status: Literal["entwurf", "in_pruefung", "freigegeben", "versendet", "abgeschlossen", "archiviert"]

class TemplateCreate(BaseModel):
    name: str
    description: str = ""
    document_type: str
    deal_types: List[Literal["ankauf", "verkauf"]] = ["ankauf", "verkauf"]
    content: str = ""
    placeholders: List[str] = []
    active: bool = True

class TemplateResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    document_type: str
    deal_types: List[str]
    placeholders: List[str]
    active: bool
    created_at: str

class AttachmentCreate(BaseModel):
    deal_id: str
    attachment_type: str
    description: str = ""
    show_in_documents: bool = True
    include_in_zip: bool = True

class AttachmentResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    deal_id: str
    attachment_type: str
    description: str
    filename: str
    show_in_documents: bool
    include_in_zip: bool
    created_at: str

class DocumentResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    deal_id: str
    template_id: str
    template_name: str
    version: int
    filename: str
    created_at: str
    created_by: str

class ChatMessage(BaseModel):
    message: str
    deal_id: Optional[str] = None

class ChatResponse(BaseModel):
    response: str
    suggestions: List[str] = []

# ==================== AUTH HELPERS ====================

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str, role: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "role": role,
        "exp": datetime.now(timezone.utc).timestamp() + (JWT_EXPIRATION_HOURS * 3600)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="Benutzer nicht gefunden")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token abgelaufen")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Ungültiger Token")

# ==================== NUMBER SERIES ====================

async def get_next_deal_number(deal_type: str) -> str:
    year = datetime.now().year
    counter_key = f"deal_{year}"
    
    counter = await db.counters.find_one_and_update(
        {"key": counter_key},
        {"$inc": {"value": 1}},
        upsert=True,
        return_document=True
    )
    
    number = counter.get("value", 1)
    return f"{year}-{str(number).zfill(3)}"

# ==================== AUTH ROUTES ====================

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user: UserCreate):
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="E-Mail bereits registriert")
    
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": user.email,
        "password": hash_password(user.password),
        "name": user.name,
        "role": user.role,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user.email, user.role)
    user_response = UserResponse(
        id=user_id,
        email=user.email,
        name=user.name,
        role=user.role,
        created_at=user_doc["created_at"]
    )
    
    return TokenResponse(access_token=token, user=user_response)

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Ungültige Anmeldedaten")
    
    token = create_token(user["id"], user["email"], user["role"])
    user_response = UserResponse(
        id=user["id"],
        email=user["email"],
        name=user["name"],
        role=user["role"],
        created_at=user["created_at"]
    )
    
    return TokenResponse(access_token=token, user=user_response)

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return UserResponse(
        id=current_user["id"],
        email=current_user["email"],
        name=current_user["name"],
        role=current_user["role"],
        created_at=current_user["created_at"]
    )

# ==================== USER MANAGEMENT ROUTES (Admin only) ====================

@api_router.get("/users", response_model=List[UserResponse])
async def get_users(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer verwalten")
    
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(100)
    return [UserResponse(**u) for u in users]

@api_router.post("/users", response_model=UserResponse)
async def create_user(user: UserCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer erstellen")
    
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="E-Mail bereits registriert")
    
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "email": user.email,
        "password": hash_password(user.password),
        "name": user.name,
        "role": user.role,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    return UserResponse(
        id=user_id,
        email=user.email,
        name=user.name,
        role=user.role,
        created_at=user_doc["created_at"]
    )

@api_router.delete("/users/{user_id}")
async def delete_user(user_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Benutzer löschen")
    
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Sie können sich nicht selbst löschen")
    
    result = await db.users.delete_one({"id": user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Benutzer nicht gefunden")
    
    return {"message": "Benutzer gelöscht"}

# ==================== DEAL ROUTES ====================

@api_router.post("/deals", response_model=DealResponse)
async def create_deal(deal: DealCreate, current_user: dict = Depends(get_current_user)):
    deal_id = str(uuid.uuid4())
    deal_number = await get_next_deal_number(deal.deal_type)
    now = datetime.now(timezone.utc).isoformat()
    
    deal_doc = {
        "id": deal_id,
        "deal_number": deal_number,
        "deal_type": deal.deal_type,
        "status": "entwurf",
        "company": deal.company.model_dump(),
        "sellers": [s.model_dump() for s in deal.sellers],
        "buyers": [b.model_dump() for b in deal.buyers],
        "vr_members": [v.model_dump() for v in deal.vr_members],
        "kaufpreis": deal.kaufpreis,
        "kaufpreis_regelung": deal.kaufpreis_regelung,
        "besitzantritt": deal.besitzantritt,
        "unterschriftsort": deal.unterschriftsort,
        "unterschriftsdatum": deal.unterschriftsdatum,
        "anzahlung": deal.anzahlung,
        "anzahlung_aktiviert": deal.anzahlung_aktiviert,
        "darlehen_uebernahme": deal.darlehen_uebernahme,
        "darlehen_betrag": deal.darlehen_betrag,
        "notizen": deal.notizen,
        "created_at": now,
        "updated_at": now,
        "created_by": current_user["id"]
    }
    
    await db.deals.insert_one(deal_doc)
    del deal_doc["_id"]
    return DealResponse(**deal_doc)

@api_router.get("/deals", response_model=List[DealResponse])
async def get_deals(
    status: Optional[str] = None,
    deal_type: Optional[str] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if status:
        query["status"] = status
    if deal_type:
        query["deal_type"] = deal_type
    if search:
        query["$or"] = [
            {"deal_number": {"$regex": search, "$options": "i"}},
            {"company.name_current": {"$regex": search, "$options": "i"}},
        ]
    
    deals = await db.deals.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return [DealResponse(**d) for d in deals]

@api_router.get("/deals/{deal_id}", response_model=DealResponse)
async def get_deal(deal_id: str, current_user: dict = Depends(get_current_user)):
    deal = await db.deals.find_one({"id": deal_id}, {"_id": 0})
    if not deal:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    return DealResponse(**deal)

@api_router.put("/deals/{deal_id}", response_model=DealResponse)
async def update_deal(deal_id: str, deal: DealCreate, current_user: dict = Depends(get_current_user)):
    existing = await db.deals.find_one({"id": deal_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    
    update_doc = {
        "company": deal.company.model_dump(),
        "sellers": [s.model_dump() for s in deal.sellers],
        "buyers": [b.model_dump() for b in deal.buyers],
        "vr_members": [v.model_dump() for v in deal.vr_members],
        "kaufpreis": deal.kaufpreis,
        "kaufpreis_regelung": deal.kaufpreis_regelung,
        "besitzantritt": deal.besitzantritt,
        "unterschriftsort": deal.unterschriftsort,
        "unterschriftsdatum": deal.unterschriftsdatum,
        "anzahlung": deal.anzahlung,
        "anzahlung_aktiviert": deal.anzahlung_aktiviert,
        "darlehen_uebernahme": deal.darlehen_uebernahme,
        "darlehen_betrag": deal.darlehen_betrag,
        "notizen": deal.notizen,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.deals.update_one({"id": deal_id}, {"$set": update_doc})
    updated = await db.deals.find_one({"id": deal_id}, {"_id": 0})
    return DealResponse(**updated)

@api_router.patch("/deals/{deal_id}/status", response_model=DealResponse)
async def update_deal_status(deal_id: str, status_update: DealStatusUpdate, current_user: dict = Depends(get_current_user)):
    existing = await db.deals.find_one({"id": deal_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    
    await db.deals.update_one(
        {"id": deal_id},
        {"$set": {"status": status_update.status, "updated_at": datetime.now(timezone.utc).isoformat()}}
    )
    
    updated = await db.deals.find_one({"id": deal_id}, {"_id": 0})
    return DealResponse(**updated)

@api_router.delete("/deals/{deal_id}")
async def delete_deal(deal_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.deals.delete_one({"id": deal_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    return {"message": "Deal gelöscht"}

# ==================== TEMPLATE ROUTES ====================

@api_router.post("/templates", response_model=TemplateResponse)
async def create_template(template: TemplateCreate, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Templates erstellen")
    
    template_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    template_doc = {
        "id": template_id,
        "name": template.name,
        "description": template.description,
        "document_type": template.document_type,
        "deal_types": template.deal_types,
        "content": template.content,
        "placeholders": template.placeholders,
        "active": template.active,
        "created_at": now
    }
    
    await db.templates.insert_one(template_doc)
    del template_doc["_id"]
    del template_doc["content"]
    return TemplateResponse(**template_doc)

@api_router.get("/templates", response_model=List[TemplateResponse])
async def get_templates(deal_type: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {"active": True}
    if deal_type:
        query["deal_types"] = deal_type
    
    templates = await db.templates.find(query, {"_id": 0, "content": 0}).to_list(100)
    return [TemplateResponse(**t) for t in templates]

@api_router.get("/templates/{template_id}")
async def get_template(template_id: str, current_user: dict = Depends(get_current_user)):
    template = await db.templates.find_one({"id": template_id}, {"_id": 0})
    if not template:
        raise HTTPException(status_code=404, detail="Template nicht gefunden")
    return template

@api_router.delete("/templates/{template_id}")
async def delete_template(template_id: str, current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Nur Administratoren können Templates löschen")
    
    result = await db.templates.delete_one({"id": template_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Template nicht gefunden")
    return {"message": "Template gelöscht"}

# ==================== DOCUMENT GENERATION ====================

def replace_placeholders(text: str, deal: dict) -> str:
    """Replace placeholders in text with deal data"""
    replacements = {
        "{{deal.number}}": deal.get("deal_number", ""),
        "{{deal.type}}": "Ankauf" if deal.get("deal_type") == "ankauf" else "Verkauf",
        "{{deal.date}}": deal.get("unterschriftsdatum", datetime.now().strftime("%d.%m.%Y")),
        "{{deal.ort}}": deal.get("unterschriftsort", "Zürich"),
        "{{company.name_current}}": deal.get("company", {}).get("name_current", ""),
        "{{company.name_new}}": deal.get("company", {}).get("name_new", "") or deal.get("company", {}).get("name_current", ""),
        "{{company.sitz_current}}": deal.get("company", {}).get("sitz_current", ""),
        "{{company.sitz_new}}": deal.get("company", {}).get("sitz_new", "") or deal.get("company", {}).get("sitz_current", ""),
        "{{company.zweck}}": deal.get("company", {}).get("zweck_current", ""),
        "{{company.che_nummer}}": deal.get("company", {}).get("che_nummer", ""),
        "{{company.hr_nummer}}": deal.get("company", {}).get("hr_nummer", ""),
        "{{company.aktienkapital}}": f"CHF {deal.get('company', {}).get('aktienkapital', 0):,.2f}",
        "{{company.anzahl_aktien}}": str(deal.get("company", {}).get("anzahl_aktien", 0)),
        "{{company.nennwert}}": f"CHF {deal.get('company', {}).get('nennwert', 0):,.2f}",
        "{{kaufpreis}}": f"CHF {deal.get('kaufpreis', 0):,.2f}",
        "{{kaufpreis_regelung}}": deal.get("kaufpreis_regelung", ""),
        "{{besitzantritt}}": deal.get("besitzantritt", ""),
        "{{anzahlung}}": f"CHF {deal.get('anzahlung', 0):,.2f}",
        "{{darlehen_betrag}}": f"CHF {deal.get('darlehen_betrag', 0):,.2f}",
    }
    
    # Build seller list
    sellers = deal.get("sellers", [])
    seller_list = "\n".join([f"- {s.get('name', '')} ({s.get('ort', '')})" for s in sellers])
    replacements["{{seller.list}}"] = seller_list or "–"
    
    # Build buyer list
    buyers = deal.get("buyers", [])
    buyer_list = "\n".join([f"- {b.get('name', '')} ({b.get('ort', '')})" for b in buyers])
    replacements["{{buyer.list}}"] = buyer_list or "–"
    
    # Build VR list
    vr_members = deal.get("vr_members", [])
    vr_list = "\n".join([f"- {v.get('name', '')} ({v.get('zeichnungsart', '')})" for v in vr_members])
    replacements["{{vr.list}}"] = vr_list or "–"
    
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, str(value))
    
    return text

def generate_docx(deal: dict, template: dict) -> io.BytesIO:
    """Generate a DOCX document from template and deal data"""
    doc = Document()
    
    # Add header with company branding
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Blum Verwaltungs- und Treuhand AG"
    
    # Add title
    doc.add_heading(template.get("name", "Dokument"), level=1)
    
    # Add deal info
    doc.add_paragraph(f"Deal-Nummer: {deal.get('deal_number', '')}")
    doc.add_paragraph(f"Datum: {deal.get('unterschriftsdatum', datetime.now().strftime('%d.%m.%Y'))}")
    doc.add_paragraph()
    
    # Add content
    content = template.get("content", "")
    if content:
        content = replace_placeholders(content, deal)
        for para in content.split("\n"):
            doc.add_paragraph(para)
    else:
        # Generate default content based on document type
        doc_type = template.get("document_type", "")
        
        if doc_type == "kaufvertrag":
            doc.add_heading("Kaufvertrag", level=2)
            doc.add_paragraph(f"Gesellschaft: {deal.get('company', {}).get('name_current', '')}")
            doc.add_paragraph(f"CHE-Nummer: {deal.get('company', {}).get('che_nummer', '')}")
            doc.add_paragraph()
            
            doc.add_heading("Verkäufer", level=3)
            for seller in deal.get("sellers", []):
                doc.add_paragraph(f"• {seller.get('name', '')} - {seller.get('address', '')} {seller.get('plz', '')} {seller.get('ort', '')}")
            
            doc.add_heading("Käufer", level=3)
            for buyer in deal.get("buyers", []):
                doc.add_paragraph(f"• {buyer.get('name', '')} - {buyer.get('address', '')} {buyer.get('plz', '')} {buyer.get('ort', '')}")
            
            doc.add_paragraph()
            doc.add_paragraph(f"Kaufpreis: CHF {deal.get('kaufpreis', 0):,.2f}")
            doc.add_paragraph(f"Besitzantritt: {deal.get('besitzantritt', '')}")
            
        elif doc_type == "quittung":
            doc.add_heading("Quittung", level=2)
            doc.add_paragraph(f"Betrag: CHF {deal.get('kaufpreis', 0):,.2f}")
            doc.add_paragraph(f"Für: {deal.get('company', {}).get('name_current', '')}")
            
        else:
            doc.add_paragraph("Dokument erstellt für Deal " + deal.get("deal_number", ""))
    
    # Add signature section
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading("Unterschriften", level=2)
    
    # Blum signature (always one)
    doc.add_paragraph()
    doc.add_paragraph("Blum Verwaltungs- und Treuhand AG")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph()
    
    # Seller signatures
    for seller in deal.get("sellers", []):
        if seller.get("requires_signature", True):
            doc.add_paragraph(f"Verkäufer: {seller.get('name', '')}")
            doc.add_paragraph("_" * 40)
            doc.add_paragraph()
    
    # Buyer signatures
    for buyer in deal.get("buyers", []):
        if buyer.get("requires_signature", True):
            doc.add_paragraph(f"Käufer: {buyer.get('name', '')}")
            doc.add_paragraph("_" * 40)
            doc.add_paragraph()
    
    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Ort/Datum: {deal.get('unterschriftsort', 'Zürich')}, {deal.get('unterschriftsdatum', datetime.now().strftime('%d.%m.%Y'))}"
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@api_router.post("/deals/{deal_id}/documents/generate/{template_id}", response_model=DocumentResponse)
async def generate_document(deal_id: str, template_id: str, current_user: dict = Depends(get_current_user)):
    deal = await db.deals.find_one({"id": deal_id}, {"_id": 0})
    if not deal:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    
    template = await db.templates.find_one({"id": template_id}, {"_id": 0})
    if not template:
        raise HTTPException(status_code=404, detail="Template nicht gefunden")
    
    # Get next version number
    existing_docs = await db.documents.count_documents({"deal_id": deal_id, "template_id": template_id})
    version = existing_docs + 1
    
    # Generate filename
    company_name = deal.get("company", {}).get("name_current", "AG").replace(" ", "_")
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"Blum_{deal.get('deal_type', 'deal').capitalize()}_{deal.get('deal_number', '')}_{template.get('document_type', 'doc')}_{company_name}_{date_str}_v{version}.docx"
    
    # Generate document
    docx_buffer = generate_docx(deal, template)
    
    # Save document info to DB
    doc_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    doc_record = {
        "id": doc_id,
        "deal_id": deal_id,
        "template_id": template_id,
        "template_name": template.get("name", ""),
        "version": version,
        "filename": filename,
        "content": docx_buffer.getvalue(),
        "created_at": now,
        "created_by": current_user["id"]
    }
    
    await db.documents.insert_one(doc_record)
    
    return DocumentResponse(
        id=doc_id,
        deal_id=deal_id,
        template_id=template_id,
        template_name=template.get("name", ""),
        version=version,
        filename=filename,
        created_at=now,
        created_by=current_user["id"]
    )

@api_router.get("/deals/{deal_id}/documents", response_model=List[DocumentResponse])
async def get_deal_documents(deal_id: str, current_user: dict = Depends(get_current_user)):
    documents = await db.documents.find({"deal_id": deal_id}, {"_id": 0, "content": 0}).sort("created_at", -1).to_list(100)
    return [DocumentResponse(**d) for d in documents]

@api_router.get("/documents/{document_id}/download")
async def download_document(document_id: str, current_user: dict = Depends(get_current_user)):
    doc = await db.documents.find_one({"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    
    content = doc.get("content")
    filename = doc.get("filename", "document.docx")
    
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ==================== ATTACHMENT ROUTES ====================

@api_router.post("/deals/{deal_id}/attachments")
async def upload_attachment(
    deal_id: str,
    file: UploadFile = File(...),
    attachment_type: str = "sonstige",
    description: str = "",
    show_in_documents: bool = True,
    include_in_zip: bool = True,
    current_user: dict = Depends(get_current_user)
):
    deal = await db.deals.find_one({"id": deal_id})
    if not deal:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    
    content = await file.read()
    attachment_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    attachment_doc = {
        "id": attachment_id,
        "deal_id": deal_id,
        "attachment_type": attachment_type,
        "description": description,
        "filename": file.filename,
        "content_type": file.content_type,
        "content": content,
        "show_in_documents": show_in_documents,
        "include_in_zip": include_in_zip,
        "created_at": now
    }
    
    await db.attachments.insert_one(attachment_doc)
    
    return AttachmentResponse(
        id=attachment_id,
        deal_id=deal_id,
        attachment_type=attachment_type,
        description=description,
        filename=file.filename,
        show_in_documents=show_in_documents,
        include_in_zip=include_in_zip,
        created_at=now
    )

@api_router.get("/deals/{deal_id}/attachments", response_model=List[AttachmentResponse])
async def get_deal_attachments(deal_id: str, current_user: dict = Depends(get_current_user)):
    attachments = await db.attachments.find({"deal_id": deal_id}, {"_id": 0, "content": 0}).to_list(100)
    return [AttachmentResponse(**a) for a in attachments]

@api_router.delete("/attachments/{attachment_id}")
async def delete_attachment(attachment_id: str, current_user: dict = Depends(get_current_user)):
    result = await db.attachments.delete_one({"id": attachment_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Anhang nicht gefunden")
    return {"message": "Anhang gelöscht"}

# ==================== CHATBOT ROUTES ====================

# Store chat sessions in memory (for simplicity)
chat_sessions = {}

SYSTEM_PROMPT = """Du bist ein hilfreicher Assistent für die Blum Verwaltungs- und Treuhand AG.
Du beantwortest Fragen zu:
- Ankauf und Verkauf von Aktiengesellschaften (AG)
- Benötigte Dokumente und Unterlagen
- Ablauf von Transaktionen
- Anforderungen an Käufer und Verkäufer
- Unterschriften und Formalitäten

Wichtige Informationen:
- Bei einem Ankauf kauft Blum eine AG
- Bei einem Verkauf verkauft Blum eine AG
- Für jeden Deal werden verschiedene Dokumente benötigt: Kaufvertrag, Quittung, Mandatsvertrag, etc.
- Verkäufer und Käufer können Personen oder Firmen sein
- Die Unterschriftenlogik: Blum hat immer einen Unterschriftsblock, Käufer/Verkäufer haben je nach Anzahl mehrere Blöcke

Antworte immer auf Deutsch und sei freundlich und professionell."""

@api_router.post("/chat", response_model=ChatResponse)
async def chat(message: ChatMessage, current_user: dict = Depends(get_current_user)):
    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="OpenAI API Key nicht konfiguriert")
    
    session_id = f"user_{current_user['id']}"
    
    try:
        # Create or get chat instance
        if session_id not in chat_sessions:
            chat_sessions[session_id] = LlmChat(
                api_key=OPENAI_API_KEY,
                session_id=session_id,
                system_message=SYSTEM_PROMPT
            ).with_model("openai", "gpt-4o")
        
        chat = chat_sessions[session_id]
        
        # Add deal context if provided
        context = message.message
        if message.deal_id:
            deal = await db.deals.find_one({"id": message.deal_id}, {"_id": 0})
            if deal:
                context = f"[Kontext: Deal {deal.get('deal_number')} - {deal.get('company', {}).get('name_current', '')} - Status: {deal.get('status')}]\n\n{message.message}"
        
        # Send message
        user_msg = UserMessage(text=context)
        response = await chat.send_message(user_msg)
        
        # Generate suggestions
        suggestions = [
            "Welche Dokumente benötige ich?",
            "Wie läuft ein Ankauf ab?",
            "Welche Informationen muss ich angeben?"
        ]
        
        return ChatResponse(response=response, suggestions=suggestions)
        
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return ChatResponse(
            response="Entschuldigung, es gab einen Fehler bei der Verarbeitung Ihrer Anfrage. Bitte versuchen Sie es später erneut.",
            suggestions=["Welche Dokumente benötige ich?", "Wie kann ich einen Deal anlegen?"]
        )

# ==================== STATS ROUTES ====================

@api_router.get("/stats")
async def get_stats(current_user: dict = Depends(get_current_user)):
    total_deals = await db.deals.count_documents({})
    active_deals = await db.deals.count_documents({"status": {"$nin": ["archiviert", "abgeschlossen"]}})
    ankauf_deals = await db.deals.count_documents({"deal_type": "ankauf"})
    verkauf_deals = await db.deals.count_documents({"deal_type": "verkauf"})
    
    # Get status counts
    pipeline = [
        {"$group": {"_id": "$status", "count": {"$sum": 1}}}
    ]
    status_counts = await db.deals.aggregate(pipeline).to_list(10)
    status_dict = {s["_id"]: s["count"] for s in status_counts}
    
    return {
        "total_deals": total_deals,
        "active_deals": active_deals,
        "ankauf_deals": ankauf_deals,
        "verkauf_deals": verkauf_deals,
        "status_counts": status_dict
    }

# ==================== SEED DEFAULT TEMPLATES ====================

async def seed_default_templates():
    """Seed default templates if none exist"""
    count = await db.templates.count_documents({})
    if count == 0:
        default_templates = [
            {
                "id": str(uuid.uuid4()),
                "name": "Kaufvertrag",
                "description": "Standard-Kaufvertrag für AG-Transaktionen",
                "document_type": "kaufvertrag",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{company.name_current}}", "{{kaufpreis}}", "{{seller.list}}", "{{buyer.list}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "name": "Quittung",
                "description": "Zahlungsquittung",
                "document_type": "quittung",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{kaufpreis}}", "{{deal.date}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "name": "Mandatsvertrag",
                "description": "Mandatsvertrag für die Treuhanddienstleistung",
                "document_type": "mandatsvertrag",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{company.name_current}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "name": "Auftragsbestätigung",
                "description": "Bestätigung des Auftrags",
                "document_type": "auftragsbestaetigung",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{company.name_current}}", "{{deal.date}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "name": "Übergabebestätigung",
                "description": "Bestätigung der Übergabe",
                "document_type": "uebergabebestaetigung",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{company.name_current}}", "{{besitzantritt}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            },
            {
                "id": str(uuid.uuid4()),
                "name": "Checkliste Abschluss",
                "description": "Checkliste für den Transaktionsabschluss",
                "document_type": "checkliste",
                "deal_types": ["ankauf", "verkauf"],
                "content": "",
                "placeholders": ["{{deal.number}}", "{{company.name_current}}"],
                "active": True,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
        ]
        
        await db.templates.insert_many(default_templates)
        logger.info("Default templates seeded")

# ==================== DEAL VALIDATION ====================

@api_router.get("/deals/{deal_id}/validate")
async def validate_deal_endpoint(deal_id: str, current_user: dict = Depends(get_current_user)):
    deal = await db.deals.find_one({"id": deal_id}, {"_id": 0})
    if not deal:
        raise HTTPException(status_code=404, detail="Deal nicht gefunden")
    
    validation_errors = []
    warnings = []
    
    # Check required fields
    if not deal.get("deal_type"):
        validation_errors.append({"field": "deal_type", "message": "Deal-Typ erforderlich"})
    
    company = deal.get("company", {})
    if not company.get("name_current"):
        validation_errors.append({"field": "company.name_current", "message": "Firmenname erforderlich"})
    
    sellers = deal.get("sellers", [])
    if not sellers:
        validation_errors.append({"field": "sellers", "message": "Mindestens ein Verkäufer erforderlich"})
    
    buyers = deal.get("buyers", [])
    if not buyers:
        validation_errors.append({"field": "buyers", "message": "Mindestens ein Käufer erforderlich"})
    
    if not deal.get("kaufpreis") and not deal.get("kaufpreis_regelung"):
        validation_errors.append({"field": "kaufpreis", "message": "Kaufpreis oder Kaufpreis-Regelung erforderlich"})
    
    if not deal.get("besitzantritt"):
        warnings.append({"field": "besitzantritt", "message": "Besitzantritt nicht definiert"})
    
    # Check signature setup
    has_signature = False
    for seller in sellers:
        if seller.get("requires_signature"):
            has_signature = True
            break
    for buyer in buyers:
        if buyer.get("requires_signature"):
            has_signature = True
            break
    
    if not has_signature:
        warnings.append({"field": "signatures", "message": "Keine Unterschriften konfiguriert"})
    
    return {
        "valid": len(validation_errors) == 0,
        "errors": validation_errors,
        "warnings": warnings
    }

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup_event():
    await seed_default_templates()

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
